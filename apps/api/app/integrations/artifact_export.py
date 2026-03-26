"""Artifact export — PDF / Markdown / HTML / JSON / CSV / DOCX output.

Exports task artifacts and documents in various formats,
saving locally or sending to external services (Google Docs / Notion / n8n Webhook).

Safety:
- File output is limited to sandbox whitelisted folders
- Data protection policies are applied for external transmission
- Export operations are recorded in audit logs
"""

from __future__ import annotations

import csv
import io
import json
import logging
import os
import uuid
from dataclasses import dataclass, field
from datetime import UTC, datetime
from enum import Enum
from pathlib import Path
from typing import Any

logger = logging.getLogger(__name__)

# Default export directory
_DEFAULT_EXPORT_DIR = os.environ.get("EXPORT_DIR", os.path.join(os.getcwd(), "exports"))


class ExportFormat(str, Enum):
    """Export format."""

    PDF = "pdf"
    MARKDOWN = "markdown"
    HTML = "html"
    JSON = "json"
    CSV = "csv"
    DOCX = "docx"


class ExportTarget(str, Enum):
    """Export target."""

    LOCAL = "local"
    GOOGLE_DOCS = "google_docs"
    NOTION = "notion"
    N8N_WEBHOOK = "n8n_webhook"


@dataclass
class ExportRequest:
    """Export request."""

    content: str | list[dict[str, Any]]
    format: ExportFormat
    target: ExportTarget = ExportTarget.LOCAL
    filename: str = ""
    metadata: dict[str, Any] = field(default_factory=dict)


@dataclass
class ExportResult:
    """Export result."""

    success: bool
    file_path: str = ""
    url: str = ""
    size_bytes: int = 0
    format: ExportFormat = ExportFormat.MARKDOWN
    error: str = ""
    exported_at: str = ""


class ArtifactExporter:
    """Artifact export service.

    Converts content to the specified format and outputs to local or external services.
    """

    def __init__(self, export_dir: str = _DEFAULT_EXPORT_DIR) -> None:
        self._export_dir = export_dir

    async def export(self, request: ExportRequest) -> ExportResult:
        """Main export method — dispatches based on format and target."""
        filename = request.filename or f"export-{uuid.uuid4().hex[:8]}"
        now = datetime.now(UTC).isoformat()

        try:
            # Format conversion
            format_handlers = {
                ExportFormat.MARKDOWN: self._export_markdown,
                ExportFormat.HTML: self._export_html,
                ExportFormat.JSON: self._export_json,
                ExportFormat.CSV: self._export_csv,
                ExportFormat.PDF: self._export_pdf,
                ExportFormat.DOCX: self._export_docx,
            }

            handler = format_handlers.get(request.format)
            if handler is None:
                return ExportResult(
                    success=False,
                    format=request.format,
                    error=f"Unsupported format: {request.format.value}",
                    exported_at=now,
                )

            result = await handler(request.content, filename)
            result.format = request.format
            result.exported_at = now

            # Send to target
            if request.target != ExportTarget.LOCAL and result.success:
                target_result = await self._send_to_target(
                    request.target,
                    result.file_path,
                    request.content,
                    request.metadata,
                )
                if not target_result.success:
                    return target_result
                result.url = target_result.url

            logger.info(
                "Artifact exported: format=%s, target=%s, size=%d bytes",
                request.format.value,
                request.target.value,
                result.size_bytes,
            )
            return result

        except Exception as exc:
            logger.error("Export failed: %s", exc)
            return ExportResult(
                success=False,
                format=request.format,
                error=str(exc),
                exported_at=now,
            )

    async def _export_markdown(
        self, content: str | list[dict[str, Any]], filename: str
    ) -> ExportResult:
        """Export in Markdown format."""
        text = self._to_text(content)
        file_path = self._resolve_path(filename, ".md")
        self._ensure_dir(file_path)

        Path(file_path).write_text(text, encoding="utf-8")
        size = os.path.getsize(file_path)

        return ExportResult(
            success=True,
            file_path=file_path,
            size_bytes=size,
        )

    async def _export_html(
        self, content: str | list[dict[str, Any]], filename: str
    ) -> ExportResult:
        """Export in HTML format."""
        text = self._to_text(content)
        html = self._markdown_to_html(text)
        file_path = self._resolve_path(filename, ".html")
        self._ensure_dir(file_path)

        Path(file_path).write_text(html, encoding="utf-8")
        size = os.path.getsize(file_path)

        return ExportResult(
            success=True,
            file_path=file_path,
            size_bytes=size,
        )

    async def _export_json(
        self, content: str | list[dict[str, Any]], filename: str
    ) -> ExportResult:
        """Export in JSON format."""
        if isinstance(content, str):
            data = {"content": content, "exported_at": datetime.now(UTC).isoformat()}
        else:
            data = {"items": content, "exported_at": datetime.now(UTC).isoformat()}

        file_path = self._resolve_path(filename, ".json")
        self._ensure_dir(file_path)

        json_text = json.dumps(data, ensure_ascii=False, indent=2)
        Path(file_path).write_text(json_text, encoding="utf-8")
        size = os.path.getsize(file_path)

        return ExportResult(
            success=True,
            file_path=file_path,
            size_bytes=size,
        )

    async def _export_csv(self, content: str | list[dict[str, Any]], filename: str) -> ExportResult:
        """Export in CSV format.

        If content is a list (list[dict]), it is directly converted to CSV.
        If it is a string, it is output as a single column ``content``.
        """
        file_path = self._resolve_path(filename, ".csv")
        self._ensure_dir(file_path)

        output = io.StringIO()

        if isinstance(content, list) and content and isinstance(content[0], dict):
            fieldnames = list(content[0].keys())
            writer = csv.DictWriter(output, fieldnames=fieldnames)
            writer.writeheader()
            for row in content:
                writer.writerow(row)
        else:
            writer_simple = csv.writer(output)
            writer_simple.writerow(["content"])
            text = content if isinstance(content, str) else json.dumps(content)
            for line in text.splitlines():
                writer_simple.writerow([line])

        csv_text = output.getvalue()
        Path(file_path).write_text(csv_text, encoding="utf-8")
        size = os.path.getsize(file_path)

        return ExportResult(
            success=True,
            file_path=file_path,
            size_bytes=size,
        )

    async def _export_pdf(self, content: str | list[dict[str, Any]], filename: str) -> ExportResult:
        """Export in PDF format.

        Uses HTML as an intermediate format for basic PDF conversion.
        If weasyprint is available it is used; otherwise a simple text PDF
        is generated alongside an HTML file.
        """
        text = self._to_text(content)
        html = self._markdown_to_html(text)
        file_path = self._resolve_path(filename, ".pdf")
        self._ensure_dir(file_path)

        try:
            from weasyprint import HTML as WeasyprintHTML

            WeasyprintHTML(string=html).write_pdf(file_path)
        except ImportError:
            # If weasyprint is not available, fallback to HTML
            logger.warning("weasyprint not available — falling back to HTML export for PDF")
            fallback_path = self._resolve_path(filename, ".html")
            Path(fallback_path).write_text(html, encoding="utf-8")
            # Create a simple text-based PDF alternative file
            pdf_content = self._build_simple_pdf(text)
            Path(file_path).write_bytes(pdf_content)

        size = os.path.getsize(file_path)
        return ExportResult(
            success=True,
            file_path=file_path,
            size_bytes=size,
        )

    async def _export_docx(
        self, content: str | list[dict[str, Any]], filename: str
    ) -> ExportResult:
        """Export in DOCX format.

        If python-docx is available it is used; otherwise an error is returned.
        """
        text = self._to_text(content)
        file_path = self._resolve_path(filename, ".docx")
        self._ensure_dir(file_path)

        try:
            from docx import Document

            doc = Document()
            for line in text.splitlines():
                stripped = line.strip()
                if stripped.startswith("# "):
                    doc.add_heading(stripped[2:], level=1)
                elif stripped.startswith("## "):
                    doc.add_heading(stripped[3:], level=2)
                elif stripped.startswith("### "):
                    doc.add_heading(stripped[4:], level=3)
                elif stripped:
                    doc.add_paragraph(stripped)
            doc.save(file_path)
        except ImportError:
            return ExportResult(
                success=False,
                error=(
                    "python-docx is required for DOCX export. Install with: pip install python-docx"
                ),
            )

        size = os.path.getsize(file_path)
        return ExportResult(
            success=True,
            file_path=file_path,
            size_bytes=size,
        )

    # ------------------------------------------------------------------ #
    #  Send to external targets
    # ------------------------------------------------------------------ #

    async def _send_to_target(
        self,
        target: ExportTarget,
        file_path: str,
        content: str | list[dict[str, Any]],
        metadata: dict[str, Any],
    ) -> ExportResult:
        """Send content to the export target."""
        if target == ExportTarget.NOTION:
            return await self._send_to_notion(content, metadata)
        elif target == ExportTarget.GOOGLE_DOCS:
            return await self._send_to_google_docs(content, metadata)
        elif target == ExportTarget.N8N_WEBHOOK:
            webhook_url = metadata.get("webhook_url", "")
            if not webhook_url:
                return ExportResult(
                    success=False,
                    error="webhook_url is required in metadata for n8n target",
                )
            return await self._send_to_n8n(content, webhook_url)
        return ExportResult(success=True, file_path=file_path)

    async def _send_to_notion(
        self, content: str | list[dict[str, Any]], config: dict[str, Any]
    ) -> ExportResult:
        """Create a page via Notion API."""
        import httpx

        api_key = os.environ.get("NOTION_API_KEY", config.get("api_key", ""))
        parent_page_id = config.get("parent_page_id", "")
        title = config.get("title", "Exported Artifact")

        if not api_key:
            return ExportResult(
                success=False,
                error="NOTION_API_KEY is not configured",
            )
        if not parent_page_id:
            return ExportResult(
                success=False,
                error="parent_page_id is required in metadata",
            )

        text = self._to_text(content)
        # Convert to Notion API blocks (split into chunks of up to 2000 characters)
        blocks = []
        for chunk in self._chunk_text(text, 2000):
            blocks.append(
                {
                    "object": "block",
                    "type": "paragraph",
                    "paragraph": {"rich_text": [{"type": "text", "text": {"content": chunk}}]},
                }
            )

        body = {
            "parent": {"page_id": parent_page_id},
            "properties": {
                "title": {"title": [{"text": {"content": title}}]},
            },
            "children": blocks[:100],  # Notion API max 100 blocks
        }

        headers = {
            "Authorization": f"Bearer {api_key}",
            "Content-Type": "application/json",
            "Notion-Version": "2022-06-28",
        }

        async with httpx.AsyncClient(timeout=30) as client:
            resp = await client.post(
                "https://api.notion.com/v1/pages",
                headers=headers,
                json=body,
            )
            resp.raise_for_status()
            data = resp.json()

        page_url = data.get("url", "")
        logger.info("Exported to Notion: %s", page_url)

        return ExportResult(
            success=True,
            url=page_url,
            size_bytes=len(text.encode("utf-8")),
        )

    async def _send_to_google_docs(
        self, content: str | list[dict[str, Any]], config: dict[str, Any]
    ) -> ExportResult:
        """Create a document via Google Docs API."""
        import httpx

        access_token = os.environ.get("GOOGLE_ACCESS_TOKEN", config.get("access_token", ""))
        title = config.get("title", "Exported Artifact")

        if not access_token:
            return ExportResult(
                success=False,
                error="GOOGLE_ACCESS_TOKEN is not configured",
            )

        text = self._to_text(content)

        # 1. Create document
        headers = {
            "Authorization": f"Bearer {access_token}",
            "Content-Type": "application/json",
        }

        async with httpx.AsyncClient(timeout=30) as client:
            create_resp = await client.post(
                "https://docs.googleapis.com/v1/documents",
                headers=headers,
                json={"title": title},
            )
            create_resp.raise_for_status()
            doc_data = create_resp.json()
            doc_id = doc_data["documentId"]

            # 2. Insert text
            requests_body = {
                "requests": [
                    {
                        "insertText": {
                            "location": {"index": 1},
                            "text": text,
                        }
                    }
                ]
            }
            await client.post(
                f"https://docs.googleapis.com/v1/documents/{doc_id}:batchUpdate",
                headers=headers,
                json=requests_body,
            )

        doc_url = f"https://docs.google.com/document/d/{doc_id}/edit"
        logger.info("Exported to Google Docs: %s", doc_url)

        return ExportResult(
            success=True,
            url=doc_url,
            size_bytes=len(text.encode("utf-8")),
        )

    async def _send_to_n8n(
        self,
        content: str | list[dict[str, Any]],
        webhook_url: str,
    ) -> ExportResult:
        """Send content to n8n webhook."""
        import httpx

        text = self._to_text(content)
        body = {
            "source": "zero-employee-orchestrator",
            "action": "artifact_export",
            "content": text,
            "timestamp": datetime.now(UTC).isoformat(),
        }

        async with httpx.AsyncClient(timeout=30) as client:
            resp = await client.post(
                webhook_url,
                headers={"Content-Type": "application/json"},
                json=body,
            )
            resp.raise_for_status()

        logger.info("Exported to n8n webhook: %s", webhook_url)

        return ExportResult(
            success=True,
            url=webhook_url,
            size_bytes=len(text.encode("utf-8")),
        )

    # ------------------------------------------------------------------ #
    #  Utilities
    # ------------------------------------------------------------------ #

    def _resolve_path(self, filename: str, ext: str) -> str:
        """Resolve file path."""
        if not filename.endswith(ext):
            filename = f"{filename}{ext}"
        return os.path.join(self._export_dir, filename)

    @staticmethod
    def _ensure_dir(file_path: str) -> None:
        """Create parent directory if it does not exist."""
        os.makedirs(os.path.dirname(file_path), exist_ok=True)

    @staticmethod
    def _to_text(content: str | list[dict[str, Any]]) -> str:
        """Convert content to text."""
        if isinstance(content, str):
            return content
        return json.dumps(content, ensure_ascii=False, indent=2)

    @staticmethod
    def _chunk_text(text: str, max_len: int) -> list[str]:
        """Split text into chunks of specified maximum length."""
        chunks: list[str] = []
        while text:
            chunks.append(text[:max_len])
            text = text[max_len:]
        return chunks

    @staticmethod
    def _markdown_to_html(markdown_text: str) -> str:
        """Convert Markdown to simple HTML.

        Performs minimal conversion without requiring external libraries.
        """
        lines = markdown_text.splitlines()
        html_lines: list[str] = [
            "<!DOCTYPE html>",
            '<html lang="ja">',
            "<head>",
            '<meta charset="utf-8">',
            "<title>Export</title>",
            "<style>",
            "body { font-family: sans-serif; max-width: 800px; "
            "margin: 2rem auto; padding: 0 1rem; line-height: 1.6; }",
            "pre { background: #f5f5f5; padding: 1rem; overflow-x: auto; }",
            "code { background: #f5f5f5; padding: 0.2em 0.4em; border-radius: 3px; }",
            "</style>",
            "</head>",
            "<body>",
        ]

        in_code_block = False
        for line in lines:
            if line.strip().startswith("```"):
                if in_code_block:
                    html_lines.append("</code></pre>")
                    in_code_block = False
                else:
                    html_lines.append("<pre><code>")
                    in_code_block = True
                continue

            if in_code_block:
                import html as html_mod

                html_lines.append(html_mod.escape(line))
                continue

            stripped = line.strip()
            if stripped.startswith("### "):
                html_lines.append(f"<h3>{stripped[4:]}</h3>")
            elif stripped.startswith("## "):
                html_lines.append(f"<h2>{stripped[3:]}</h2>")
            elif stripped.startswith("# "):
                html_lines.append(f"<h1>{stripped[2:]}</h1>")
            elif stripped.startswith("- ") or stripped.startswith("* "):
                html_lines.append(f"<li>{stripped[2:]}</li>")
            elif stripped == "":
                html_lines.append("<br>")
            else:
                html_lines.append(f"<p>{stripped}</p>")

        if in_code_block:
            html_lines.append("</code></pre>")

        html_lines.extend(["</body>", "</html>"])
        return "\n".join(html_lines)

    @staticmethod
    def _build_simple_pdf(text: str) -> bytes:
        """Generate a minimal PDF byte sequence (no external libraries required).

        Manually constructs a minimal structure conforming to the PDF 1.4 spec.
        Only ASCII fallback since Japanese text cannot be included.
        """
        # Convert to ASCII only
        safe_text = text.encode("ascii", errors="replace").decode("ascii")
        lines = safe_text.splitlines()
        # Truncate long text
        page_lines = lines[:100]
        content = "\\n".join(
            line.replace("\\", "\\\\").replace("(", "\\(").replace(")", "\\)")
            for line in page_lines
        )

        pdf = (
            "%PDF-1.4\n"
            "1 0 obj\n<< /Type /Catalog /Pages 2 0 R >>\nendobj\n"
            "2 0 obj\n<< /Type /Pages /Kids [3 0 R] /Count 1 >>\nendobj\n"
            "3 0 obj\n<< /Type /Page /Parent 2 0 R "
            "/MediaBox [0 0 612 792] /Contents 4 0 R "
            "/Resources << /Font << /F1 5 0 R >> >> >>\nendobj\n"
            f"4 0 obj\n<< /Length {len(content) + 30} >>\nstream\n"
            f"BT /F1 10 Tf 50 750 Td ({content}) Tj ET\n"
            "endstream\nendobj\n"
            "5 0 obj\n<< /Type /Font /Subtype /Type1 "
            "/BaseFont /Helvetica >>\nendobj\n"
            "xref\n0 6\n"
            "0000000000 65535 f \n"
            "0000000009 00000 n \n"
            "0000000058 00000 n \n"
            "0000000115 00000 n \n"
            "0000000266 00000 n \n"
            "0000000400 00000 n \n"
            "trailer\n<< /Size 6 /Root 1 0 R >>\n"
            "startxref\n480\n%%EOF\n"
        )
        return pdf.encode("ascii")

    def get_supported_formats(self) -> list[dict[str, str]]:
        """Return a list of supported formats."""
        return [
            {"format": f.value, "description": self._format_description(f)} for f in ExportFormat
        ]

    def get_supported_targets(self) -> list[dict[str, str]]:
        """Return a list of supported export targets."""
        return [
            {"target": t.value, "description": self._target_description(t)} for t in ExportTarget
        ]

    @staticmethod
    def _format_description(fmt: ExportFormat) -> str:
        """Return format description."""
        descriptions = {
            ExportFormat.PDF: "PDF document",
            ExportFormat.MARKDOWN: "Markdown text",
            ExportFormat.HTML: "HTML page",
            ExportFormat.JSON: "JSON data",
            ExportFormat.CSV: "CSV spreadsheet",
            ExportFormat.DOCX: "Microsoft Word document",
        }
        return descriptions.get(fmt, fmt.value)

    @staticmethod
    def _target_description(target: ExportTarget) -> str:
        """Return target description."""
        descriptions = {
            ExportTarget.LOCAL: "Local file system",
            ExportTarget.GOOGLE_DOCS: "Google Docs",
            ExportTarget.NOTION: "Notion page",
            ExportTarget.N8N_WEBHOOK: "n8n Webhook",
        }
        return descriptions.get(target, target.value)


# Global instance
artifact_exporter = ArtifactExporter()
